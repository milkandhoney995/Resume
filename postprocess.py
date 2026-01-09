from docx import Document
from docx.shared import Inches
from pathlib import Path
import sys

# Get language from command line argument, default to Japanese
language = sys.argv[1] if len(sys.argv) > 1 else "ja"
filename = "CV.docx" if language == "en" else "職務経歴書.docx"

path = Path(filename)
doc = Document(str(path))

for table in doc.tables:
    table.autofit = False
    col_count = len(table.columns)

    # ===== 2列テーブル（期間｜業務内容）=====
    if col_count == 2:
        left_width = Inches(2.0)
        right_width = Inches(4.7)

        for row in table.rows:
            row.cells[0].width = left_width
            row.cells[1].width = right_width

    # ===== 4列テーブル（スキル）=====
    elif col_count == 4:
        w1 = Inches(1.3)  # カテゴリ
        w2 = Inches(1.8)  # スキル名
        w3 = Inches(1.0)  # 使用期間
        w4 = Inches(3.0)  # レベル

        for row in table.rows:
            row.cells[0].width = w1
            row.cells[1].width = w2
            row.cells[2].width = w3
            row.cells[3].width = w4

    # ===== 共通：擬似箇条書き処理 =====
    for row in table.rows:
        for cell in row.cells:
            text = cell.text
            if "・" in text:
                cell.text = ""
                p = cell.paragraphs[0]
                lines = [line.strip() for line in text.split("・") if line.strip()]

                for i, line in enumerate(lines):
                    run = p.add_run("・" + line)
                    if i != len(lines) - 1:
                        run.add_break()

doc.save(str(path))